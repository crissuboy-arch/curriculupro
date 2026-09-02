import base64
import json
import os
import re
import tempfile
import time
import urllib.parse
from pathlib import Path

from config import (
    CONFIG,
    MODELS as APP_MODELS,
    get_config,
    list_configs,
    reload_config,
    PLATFORM_PATTERNS,
    PRODUCT_TYPE_KEYWORDS,
    SYSTEM_PROMPTS,
    UPLOADS,
)

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Request, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from openai import OpenAI

try:
    import requests as _requests
except ImportError:
    _requests = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.shared import WD_ALIGN_PARAGRAPH
    except ImportError:
        DocxDocument = None
        Pt = None
        RGBColor = None
        WD_ALIGN_PARAGRAPH = None

ROOT = Path(__file__).parent
load_dotenv(ROOT / ".env")

BASE_URL = "https://integrate.api.nvidia.com/v1"
DEFAULT_MODEL = APP_MODELS.get("default", "meta/llama-3.1-70b-instruct")
REQUEST_TIMEOUT = 60


def _load_skills():
    skills_dir = ROOT / "skills"
    skills = {}
    if not skills_dir.is_dir():
        return skills

    def _walk_and_load(path, prefix=""):
        for entry in sorted(path.iterdir()):
            if entry.is_file() and entry.suffix.lower() == ".md" and entry.stem == "SKILL":
                parts = prefix.split("/") if prefix else []
                model_id = "_".join(parts) if parts else entry.stem
                if not model_id.strip():
                    model_id = entry.parent.stem
                content = entry.read_text(encoding="utf-8")
                words = model_id.replace("_", " ").title().split()
                label = " ".join(words) + " AI"
                skills[model_id] = {"label": label, "content": content}
            elif entry.is_file() and entry.suffix.lower() == ".md" and entry.stem != "SKILL":
                if not prefix:
                    model_id = entry.stem
                    words = model_id.replace("-", " ").title().split()
                    label = " ".join(words) + " AI"
                    content = entry.read_text(encoding="utf-8")
                    skills[model_id] = {"label": label, "content": content}
            elif entry.is_dir():
                new_prefix = prefix + "/" + entry.name if prefix else entry.name
                _walk_and_load(entry, new_prefix)

    _walk_and_load(skills_dir)
    return skills


def _build_registry(skills=None):
    if skills is None:
        skills = {}
    default_key = os.getenv("NVIDIA_API_KEY")
    if not default_key:
        raise RuntimeError(
            "NVIDIA_API_KEY nao encontrada. Local: confira o .env. "
            "Na Vercel: importe o .env.vercel em Project Settings -> Environment Variables."
        )
    registry = {
        "meta/llama-3.1-70b-instruct": {
            "label": "Llama 3.1 70B",
            "provider": "Meta",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_LLAMA", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096},
        },
        "meta/llama-3.1-405b-instruct": {
            "label": "Llama 3.1 405B",
            "provider": "Meta",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_LLAMA405", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096},
        },
        "deepseek-ai/deepseek-r1": {
            "label": "DeepSeek R1",
            "provider": "DeepSeek",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_DEEPSEEK", default_key),
            "params": {"temperature": 0.6, "top_p": 0.95, "max_tokens": 4096},
        },
        "qwen/qwen3.5-397b-a17b": {
            "label": "Qwen 3.5 397B",
            "provider": "Qwen",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_QWEN", default_key),
            "params": {"temperature": 0.6, "top_p": 0.95, "max_tokens": 4096},
        },
        "z-ai/glm-5.2": {
            "label": "GLM-5.2",
            "provider": "Z.ai",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_GLM", default_key),
            "params": {"temperature": 1, "top_p": 1, "max_tokens": 16384},
        },
        "nvidia/nemotron-3-super-120b-instruct": {
            "label": "Nemotron 120B",
            "provider": "NVIDIA",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_NEMOTRON", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096},
        },
        "mistralai/mixtral-8x22b-instruct-v0.1": {
            "label": "Mixtral 8x22B",
            "provider": "Mistral",
            "category": "llm",
            "api_key": os.getenv("NVIDIA_API_KEY_MISTRAL", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096},
        },
        "meta/llama-3.2-90b-vision-instruct": {
            "label": "Llama 3.2 90B Vision",
            "provider": "Meta",
            "category": "vision",
            "api_key": os.getenv("NVIDIA_API_KEY_VISION", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 2048},
        },
    }

    for skill_id, skill in skills.items():
        registry[skill_id] = {
            "label": skill["label"],
            "provider": "Expert",
            "category": "llm",
            "api_key": os.getenv(f"NVIDIA_API_KEY_{skill_id.upper().replace('-', '_')}", default_key),
            "params": {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096},
            "underlying": APP_MODELS.get("expert_underlying", DEFAULT_MODEL),
            "skill_content": skill["content"],
        }

    return registry


def create_app():
    skills = _load_skills()
    model_registry = _build_registry(skills)
    default_key = os.getenv("NVIDIA_API_KEY")
    clients = {}

    def client_for(model_id):
        entry = model_registry.get(model_id)
        key = entry["api_key"] if entry else default_key
        params = entry["params"] if entry else {"temperature": 0.7, "top_p": 0.95, "max_tokens": 4096}
        if key not in clients:
            clients[key] = OpenAI(base_url=BASE_URL, api_key=key, timeout=REQUEST_TIMEOUT)
        return clients[key], params

    def _sse_stream(client, model, messages, params):
        start = time.monotonic()
        first_token_at = None
        try:
            completion = client.chat.completions.create(
                model=model, messages=messages, stream=True, **params,
            )
            for chunk in completion:
                if not getattr(chunk, "choices", None):
                    continue
                choice = chunk.choices[0]
                delta = getattr(choice, "delta", None)
                if delta is None:
                    continue
                content = getattr(delta, "content", None)
                if content:
                    if first_token_at is None:
                        first_token_at = time.monotonic()
                        ttft_ms = int((first_token_at - start) * 1000)
                        yield f"event: meta\ndata: {json.dumps({'ttft_ms': ttft_ms})}\n\n"
                    yield f"data: {json.dumps({'content': content})}\n\n"
            total_ms = int((time.monotonic() - start) * 1000)
            yield f"event: done\ndata: {json.dumps({'total_ms': total_ms})}\n\n"
        except Exception as exc:
            yield f"event: error\ndata: {json.dumps({'message': str(exc)})}\n\n"

    def _sse_response(stream_gen):
        return StreamingResponse(
            stream_gen,
            media_type="text/event-stream",
            headers={
                "X-Accel-Buffering": "no",
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
            },
        )

    def _stream_analysis(user_prompt: str, model_id: str = None):
        if model_id is None:
            model_id = CONFIG.get("app_id", "pro")
        entry = model_registry.get(model_id)
        if not entry:
            return JSONResponse({"error": "Modelo nao encontrado"}, status_code=500)

        skill_content = entry.get("skill_content", "")
        underlying = entry.get("underlying", model_id)
        client, params = client_for(model_id)

        messages = [
            {"role": "system", "content": skill_content},
            {"role": "user", "content": user_prompt},
        ]

        return _sse_response(_sse_stream(client, underlying, messages, params))

    def _detect_platform(url: str) -> str:
        for pattern, name in PLATFORM_PATTERNS:
            if re.search(pattern, url, re.IGNORECASE):
                return name
        return ""

    def _detect_product_type(text: str) -> list:
        found = []
        text_lower = text.lower()
        for ptype, keywords in PRODUCT_TYPE_KEYWORDS.items():
            for kw in keywords:
                if kw in text_lower:
                    found.append(ptype)
                    break
        return found

    def _fetch_and_clean_url(url: str) -> dict:
        if _requests is None:
            raise HTTPException(500, "Biblioteca 'requests' nao instalada. Execute: pip install requests beautifulsoup4")
        if BeautifulSoup is None:
            raise HTTPException(500, "Biblioteca 'beautifulsoup4' nao instalada. Execute: pip install beautifulsoup4")

        parsed = urllib.parse.urlparse(url)
        hostname = parsed.hostname or ""

        if parsed.scheme == "file":
            raise HTTPException(400, "URLs de arquivo nao sao permitidas")
        if hostname in ("localhost", "127.0.0.1", "0.0.0.0", "::1"):
            raise HTTPException(400, "URLs locais nao sao permitidas")
        if hostname.endswith(".local") or hostname.endswith(".internal"):
            raise HTTPException(400, "URLs internas nao sao permitidas")

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }

        try:
            resp = _requests.get(url, headers=headers, timeout=15, allow_redirects=True)
            resp.raise_for_status()
        except _requests.Timeout:
            raise HTTPException(408, "A pagina demorou muito para responder. Tente novamente.")
        except _requests.HTTPError as e:
            code = e.response.status_code
            friendly = {403: "Acesso negado (403). A pagina pode exigir login.", 404: "Pagina nao encontrada (404)."}
            raise HTTPException(code, friendly.get(code, f"Erro HTTP {code} ao acessar a pagina."))
        except _requests.ConnectionError:
            raise HTTPException(502, "Nao foi possivel conectar ao servidor. Verifique a URL.")
        except _requests.RequestException as e:
            raise HTTPException(400, f"Erro ao acessar a URL: {str(e)[:120]}")

        ct = resp.headers.get("content-type", "")
        if "text/html" not in ct and "application/xhtml" not in ct:
            raise HTTPException(400, "A URL nao retornou uma pagina HTML.")

        soup = BeautifulSoup(resp.text, "html.parser")
        title = soup.title.get_text(strip=True) if soup.title else ""
        meta_desc = ""
        meta_tag = soup.find("meta", attrs={"name": "description"})
        if meta_tag and meta_tag.get("content"):
            meta_desc = meta_tag["content"].strip()

        for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe", "form", "svg", "link", "meta", "noscript"]):
            tag.decompose()

        body = soup.find("body")
        text = body.get_text(separator="\n", strip=True) if body else ""
        text = re.sub(r"\n{4,}", "\n\n", text)
        text = text[:10000]

        return {"title": title, "description": meta_desc, "content": text}

    def _generate_cv_html(cv_data, template="europe"):
        personal = cv_data.get("personal", {})
        name = personal.get("name", "Seu Nome")
        title_prof = personal.get("title", "")
        email = personal.get("email", "")
        phone = personal.get("phone", "")
        location = personal.get("location", "")
        country = personal.get("country", "")
        nationality = personal.get("nationality", "")
        dob = personal.get("dob", "")
        linkedin = personal.get("linkedin", "")
        portfolio = personal.get("portfolio", "")
        target_job = cv_data.get("target_job", "")
        summary = cv_data.get("summary", "")
        experience = cv_data.get("experience", [])
        education = cv_data.get("education", [])
        skills_tech = cv_data.get("skills_technical", cv_data.get("skills", []))
        skills_pers = cv_data.get("skills_personal", [])
        languages = cv_data.get("languages", [])
        certifications = cv_data.get("certifications", [])
        courses = cv_data.get("courses", [])
        links = cv_data.get("links", [])
        dl = cv_data.get("driving_license", {})
        avail = cv_data.get("availability", {})
        wa = cv_data.get("work_authorization", {})

        contact_parts = []
        if email:
            contact_parts.append(email)
        if phone:
            contact_parts.append(phone)
        if location:
            contact_parts.append(location)
        if country:
            contact_parts.append(country)
        if nationality:
            contact_parts.append("Nac.: " + nationality)
        if dob:
            contact_parts.append(dob)
        contact_line = " &bull; ".join(contact_parts)

        link_parts = []
        if linkedin:
            link_parts.append("LinkedIn: " + linkedin)
        if portfolio:
            link_parts.append("Portfolio: " + portfolio)
        for lk in links:
            if lk.get("label") and lk.get("url"):
                link_parts.append(lk["label"] + ": " + lk["url"])
        link_line = " &bull; ".join(link_parts)

        if template == "ats":
            border_style = "border-bottom: 2px solid #333;"
            section_color = "#333"
            font_family = "Arial, Helvetica, sans-serif"
            name_size = "24px"
            accent = "#333"
        elif template == "executive":
            border_style = "border-bottom: 3px solid #1a1a2e;"
            section_color = "#1a1a2e"
            font_family = "Georgia, 'Times New Roman', serif"
            name_size = "26px"
            accent = "#1a1a2e"
        else:
            border_style = "border-bottom: 2px solid #2563EB;"
            section_color = "#2563EB"
            font_family = "'Inter', 'Segoe UI', sans-serif"
            name_size = "24px"
            accent = "#2563EB"

        exp_html = ""
        for exp in experience:
            if exp.get("role") or exp.get("company"):
                exp_html += '<div class="exp-item"><div class="exp-header"><span class="exp-role">' + (exp.get("role", "") or "") + '</span> | <span class="exp-company">' + (exp.get("company", "") or "") + '</span></div>'
                if exp.get("period"):
                    exp_html += '<div class="exp-period">' + exp["period"] + '</div>'
                exp_html += '<ul>'
                for resp in exp.get("responsibilities", []):
                    if resp:
                        exp_html += '<li>' + resp + '</li>'
                for res in exp.get("results", []):
                    if res:
                        exp_html += '<li>' + res + '</li>'
                exp_html += '</ul></div>'

        edu_html = ""
        for edu in education:
            if edu.get("degree") or edu.get("institution"):
                edu_html += '<div class="edu-item"><div class="edu-header"><span class="edu-degree">' + (edu.get("degree", "") or "") + '</span> | <span class="edu-inst">' + (edu.get("institution", "") or "") + '</span></div>'
                if edu.get("period"):
                    edu_html += '<div class="edu-period">' + edu["period"] + '</div>'
                if edu.get("details"):
                    edu_html += '<div class="edu-details">' + edu["details"] + '</div>'
                edu_html += '</div>'

        skills_html = ""
        tech_skills = [s for s in skills_tech if s]
        pers_skills = [s for s in skills_pers if s]
        if tech_skills:
            skills_html += '<div class="section"><div class="section-title">Technical Skills</div><div class="skills-line">' + " &bull; ".join('<span class="skill-tag">' + s + '</span>' for s in tech_skills) + '</div></div>'
        if pers_skills:
            skills_html += '<div class="section"><div class="section-title">Personal Skills</div><div class="skills-line">' + " &bull; ".join('<span class="skill-tag">' + s + '</span>' for s in pers_skills) + '</div></div>'

        lang_html = ""
        for lang in languages:
            if lang.get("language"):
                lang_html += '<div class="lang-item"><span class="lang-name">' + lang["language"] + '</span> - <span class="lang-level">' + (lang.get("level", "") or "") + '</span></div>'

        cert_html = ""
        for cert in certifications:
            if cert.get("name"):
                parts = [cert["name"]]
                if cert.get("institution"):
                    parts.append(cert["institution"])
                if cert.get("year"):
                    parts.append(cert["year"])
                if cert.get("validity"):
                    parts.append("Val: " + cert["validity"])
                if cert.get("credential"):
                    parts.append("No: " + cert["credential"])
                cert_html += '<div class="cert-item">' + " &bull; ".join(parts) + '</div>'

        course_html = ""
        for course in courses:
            if course.get("name"):
                parts = [course["name"]]
                if course.get("institution"):
                    parts.append(course["institution"])
                if course.get("period"):
                    parts.append(course["period"])
                course_html += '<div class="course-item">' + " &bull; ".join(parts) + '</div>'

        links_line_html = ""
        if link_line:
            links_line_html = '<div class="links-line">' + link_line + '</div>'

        title_line_html = ""
        if title_prof:
            title_line_html = '<div class="title-line">' + title_prof + '</div>'

        summary_section = ""
        if summary:
            summary_section = '<div class="section"><div class="section-title">Professional Summary</div><div class="summary-text">' + summary + '</div></div>'

        experience_section = ""
        if exp_html:
            experience_section = '<div class="section"><div class="section-title">Work Experience</div>' + exp_html + '</div>'

        education_section = ""
        if edu_html:
            education_section = '<div class="section"><div class="section-title">Education</div>' + edu_html + '</div>'

        skills_section = skills_html

        languages_section = ""
        if lang_html:
            languages_section = '<div class="section"><div class="section-title">Languages</div>' + lang_html + '</div>'

        certifications_section = ""
        if cert_html:
            certifications_section = '<div class="section"><div class="section-title">Certifications</div>' + cert_html + '</div>'

        courses_section = ""
        if course_html:
            courses_section = '<div class="section"><div class="section-title">Courses</div>' + course_html + '</div>'

        target_job_section = ""
        if target_job:
            target_job_section = '<div class="section" style="background:#EFF6FF;padding:6px 12px;border-radius:4px;font-size:10pt;color:#1D4ED8;font-weight:600">Target position: ' + target_job + '</div>'

        driving_section = ""
        if dl.get("has"):
            cats = [c for c in dl.get("categories", []) if c]
            dl_parts = []
            if cats:
                dl_parts.append("Categories: " + ", ".join(cats))
            if dl.get("has_vehicle"):
                dl_parts.append("Own vehicle")
            if dl_parts:
                driving_section = '<div class="section"><div class="section-title">Driving License</div><div class="cert-item">' + " &bull; ".join(dl_parts) + '</div></div>'

        avail_section = ""
        avail_items = []
        if avail.get("immediate"): avail_items.append("Immediate availability")
        if avail.get("shifts"): avail_items.append("Shifts")
        if avail.get("weekends"): avail_items.append("Weekends")
        if avail.get("relocations"): avail_items.append("Relocations")
        if avail.get("relocate"): avail_items.append("Willing to relocate")
        if avail_items:
            avail_section = '<div class="section"><div class="section-title">Availability</div><div class="cert-item">' + " &bull; ".join(avail_items) + '</div></div>'

        auth_section = ""
        if wa.get("status") or wa.get("residence_permit"):
            auth_parts = []
            if wa.get("status"): auth_parts.append("Authorization: " + wa["status"])
            if wa.get("residence_permit"): auth_parts.append(wa["residence_permit"])
            auth_section = '<div class="section"><div class="section-title">Work Authorization</div><div class="cert-item">' + " &bull; ".join(auth_parts) + '</div></div>'

        html = "<!DOCTYPE html>\n<html lang=\"pt-BR\">\n<head>\n<meta charset=\"UTF-8\">\n"
        html += '<meta name="viewport" content="width=device-width, initial-scale=1.0">\n'
        html += "<title>" + name + " - CV</title>\n"
        html += '<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Georgia&display=swap" rel="stylesheet">\n'
        html += "<style>\n"
        html += "*{margin:0;padding:0;box-sizing:border-box}\n"
        html += "@page{size:A4;margin:0}\n"
        html += "body{font-family:" + font_family + ";color:#1a1a2e;line-height:1.5;background:#fff;font-size:10pt}\n"
        html += ".page{width:210mm;min-height:297mm;margin:0 auto;padding:20mm 18mm;background:#fff}\n"
        html += ".name{font-size:" + name_size + ";font-weight:700;color:#1a1a2e;margin-bottom:2px;letter-spacing:-0.02em}\n"
        html += ".contact{font-size:9pt;color:#555;margin-bottom:4px}\n"
        html += ".links-line{font-size:8.5pt;color:#666;margin-bottom:6px}\n"
        html += ".title-line{font-size:11pt;color:" + accent + ";font-style:italic;margin-bottom:12px}\n"
        html += ".section{margin-bottom:12px}\n"
        html += ".section-title{font-size:12pt;font-weight:700;color:" + accent + ";" + border_style + "padding-bottom:3px;margin-bottom:8px;text-transform:uppercase;letter-spacing:0.05em}\n"
        html += ".summary-text{font-size:9.5pt;color:#333;line-height:1.6}\n"
        html += ".exp-item{margin-bottom:10px}\n"
        html += ".exp-header{font-size:10pt;font-weight:600;color:#1a1a2e}\n"
        html += ".exp-role{font-weight:700}\n"
        html += ".exp-company{color:#444}\n"
        html += ".exp-period{font-size:8.5pt;color:#888;margin-bottom:3px}\n"
        html += ".exp-item ul{margin-left:16px;margin-top:3px}\n"
        html += ".exp-item li{font-size:9.5pt;color:#333;margin-bottom:2px;line-height:1.5}\n"
        html += ".edu-item{margin-bottom:8px}\n"
        html += ".edu-header{font-size:10pt;font-weight:600;color:#1a1a2e}\n"
        html += ".edu-degree{font-weight:700}\n"
        html += ".edu-inst{color:#444}\n"
        html += ".edu-period{font-size:8.5pt;color:#888}\n"
        html += ".edu-details{font-size:9pt;color:#555;margin-top:2px}\n"
        html += ".skills-line{font-size:9.5pt;color:#333;line-height:1.8}\n"
        html += ".skill-tag{display:inline-block;margin-right:4px}\n"
        html += ".lang-item{font-size:9.5pt;color:#333;margin-bottom:2px}\n"
        html += ".lang-name{font-weight:600}\n"
        html += ".cert-item,.course-item{font-size:9.5pt;color:#333;margin-bottom:2px}\n"
        html += "</style>\n</head>\n<body>\n<div class=\"page\">\n"
        html += '<div class="name">' + name + '</div>\n'
        html += '<div class="contact">' + contact_line + '</div>\n'
        html += links_line_html + "\n"
        html += title_line_html + "\n"
        html += target_job_section + "\n"
        html += summary_section + "\n"
        html += experience_section + "\n"
        html += education_section + "\n"
        html += skills_section + "\n"
        html += languages_section + "\n"
        html += certifications_section + "\n"
        html += courses_section + "\n"
        html += driving_section + "\n"
        html += avail_section + "\n"
        html += auth_section + "\n"
        html += "</div>\n</body>\n</html>"
        return html

    app = FastAPI()
    static_dir = ROOT / "static"

    @app.get("/")
    def index():
        return FileResponse(static_dir / "index.html")

    @app.get("/api/models")
    def list_models():
        return JSONResponse([
            {
                "id": model_id,
                "label": entry["label"],
                "provider": entry["provider"],
                "category": entry.get("category", "llm"),
            }
            for model_id, entry in model_registry.items()
        ])

    @app.get("/api/health")
    def health():
        return JSONResponse({"status": "ok", "models": list(model_registry.keys())})

    @app.get("/api/branding")
    def branding(app_id: str = None):
        if app_id:
            try:
                cfg = get_config(app_id)
                reload_config()
                return JSONResponse(cfg)
            except FileNotFoundError:
                raise HTTPException(404, f"Aplicativo '{app_id}' nao encontrado.")
        return JSONResponse(CONFIG)

    @app.get("/api/configs")
    def available_configs():
        return JSONResponse(list_configs())

    @app.post("/api/chat")
    async def chat(request: Request):
        body = await request.json()
        messages = body.get("messages", [])
        model = (body.get("model") or DEFAULT_MODEL).strip()

        entry = model_registry.get(model, {})
        skill_content = entry.get("skill_content")
        if skill_content:
            messages = [{"role": "system", "content": skill_content}] + messages

        client, params = client_for(model)
        api_model = entry.get("underlying", model)

        return _sse_response(_sse_stream(client, api_model, messages, params))

    @app.post("/api/upload-image")
    async def upload_image(file: UploadFile = File(...)):
        if not UPLOADS.get("image", True):
            raise HTTPException(400, "Upload de imagem nao habilitado para este aplicativo.")

        valid_ext = {"jpg", "jpeg", "png", "webp"}
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if ext not in valid_ext:
            raise HTTPException(400, "Formato inv\u00e1lido. Aceitos: JPG, JPEG, PNG, WEBP")

        contents = await file.read()
        if len(contents) > 10 * 1024 * 1024:
            raise HTTPException(400, "Arquivo muito grande. M\u00e1ximo: 10MB")

        b64 = base64.b64encode(contents).decode("utf-8")
        mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp"}
        mime = mime_map.get(ext, "image/jpeg")

        system_prompt = SYSTEM_PROMPTS.get("image", "Analise esta imagem.")
        vision_key = os.getenv("NVIDIA_API_KEY_VISION", default_key)
        vision_client = OpenAI(base_url=BASE_URL, api_key=vision_key, timeout=REQUEST_TIMEOUT)
        vision_model = APP_MODELS.get("vision", "meta/llama-3.2-90b-vision-instruct")

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text", "text": "Analise esta imagem de produto ou an\u00fancio detalhadamente."},
            ]},
        ]

        vision_params = {"temperature": 0.7, "top_p": 0.95, "max_tokens": 2048}
        return _sse_response(_sse_stream(vision_client, vision_model, messages, vision_params))

    @app.post("/api/upload-pdf")
    async def upload_pdf(file: UploadFile = File(...)):
        if not UPLOADS.get("pdf", True):
            raise HTTPException(400, "Upload de PDF nao habilitado para este aplicativo.")

        if PdfReader is None:
            raise HTTPException(500, "Biblioteca 'pypdf' nao instalada. Execute: pip install pypdf")

        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(400, "Formato inv\u00e1lido. Aceito apenas PDF.")
        contents = await file.read()
        if len(contents) > 20 * 1024 * 1024:
            raise HTTPException(400, "Arquivo muito grande. M\u00e1ximo: 20MB")

        safe_name = re.sub(r"[^\w\.\-]", "_", file.filename)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        try:
            tmp.write(contents)
            tmp.close()
            reader = PdfReader(tmp.name)
            text_parts = [p.extract_text() or "" for p in reader.pages]
            full_text = "\n".join(text_parts).strip()
            if len(full_text) < 50:
                return JSONResponse({"error": "Este PDF parece ser digitalizado e nao contem texto extraivel. Tente fazer upload de uma imagem do documento."}, status_code=400)
            if len(full_text) > 15000:
                full_text = full_text[:15000] + "\n[... conteudo truncado por limite de tamanho]"
        finally:
            os.unlink(tmp.name)

        instruction = SYSTEM_PROMPTS.get("pdf", "Analise o documento.")

        return _stream_analysis(f"{instruction}\n\nConteudo extraido do PDF ({safe_name}):\n\n{full_text}")

    @app.post("/api/upload-screenshot")
    async def upload_screenshot(file: UploadFile = File(...)):
        if not UPLOADS.get("screenshot", True):
            raise HTTPException(400, "Upload de screenshot nao habilitado para este aplicativo.")

        valid_ext = {"jpg", "jpeg", "png", "webp"}
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        if ext not in valid_ext:
            raise HTTPException(400, "Formato invalido. Aceitos: JPG, JPEG, PNG, WEBP")
        contents = await file.read()
        if len(contents) > 10 * 1024 * 1024:
            raise HTTPException(400, "Arquivo muito grande. Maximo: 10MB")

        b64 = base64.b64encode(contents).decode("utf-8")
        mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "webp": "image/webp"}
        mime = mime_map.get(ext, "image/jpeg")

        system_prompt = SYSTEM_PROMPTS.get("screenshot", "")
        if not system_prompt:
            raise HTTPException(400, "Upload de screenshot nao disponivel para este aplicativo.")

        vision_key = os.getenv("NVIDIA_API_KEY_VISION", default_key)
        vision_client = OpenAI(base_url=BASE_URL, api_key=vision_key, timeout=REQUEST_TIMEOUT)
        vision_model = APP_MODELS.get("vision", "meta/llama-3.2-90b-vision-instruct")

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text", "text": "Analise esta captura de tela detalhadamente."},
            ]},
        ]

        vision_params = {"temperature": 0.7, "top_p": 0.95, "max_tokens": 2048}
        return _sse_response(_sse_stream(vision_client, vision_model, messages, vision_params))

    @app.post("/api/analyze-url")
    async def analyze_url(request: Request):
        if not UPLOADS.get("url", True):
            raise HTTPException(400, "Analise de URL nao habilitada para este aplicativo.")

        body = await request.json()
        url = body.get("url", "").strip()
        analysis_type = body.get("type", "lp")

        if not url:
            raise HTTPException(400, "URL nao informada.")
        if not url.startswith("http://") and not url.startswith("https://"):
            raise HTTPException(400, "URL deve comecar com http:// ou https://")

        page_data = _fetch_and_clean_url(url)

        platform = _detect_platform(url)
        product_types = _detect_product_type(url + " " + page_data["title"] + " " + page_data["content"][:3000])
        platform_info = f"Plataforma detectada: {platform}\n" if platform else ""
        product_info = f"Tipo de produto detectado: {', '.join(product_types)}\n" if product_types else ""

        url_prompts = SYSTEM_PROMPTS.get("url", {})
        instruction = url_prompts.get(analysis_type) or url_prompts.get("lp", "Analise esta pagina.")
        user_content = (
            f"{instruction}\n\n"
            f"URL analisada: {url}\n"
            f"{platform_info}"
            f"{product_info}"
            f"Titulo da pagina: {page_data['title']}\n"
            f"Meta Description: {page_data['description']}\n\n"
            f"Conteudo extraido:\n{page_data['content']}"
        )
        return _stream_analysis(user_content)

    @app.post("/api/embeddings")
    async def embeddings(request: Request):
        body = await request.json()
        text = body.get("input", body.get("text", ""))
        model = body.get("model", "nvidia/nv-embedqa-e5-v5")
        client_obj = OpenAI(base_url=BASE_URL, api_key=default_key, timeout=REQUEST_TIMEOUT)
        try:
            start = time.monotonic()
            result = client_obj.embeddings.create(
                model=model,
                input=[text] if isinstance(text, str) else text,
                extra_body={"input_type": "query"},
            )
            total_ms = int((time.monotonic() - start) * 1000)
            return JSONResponse({
                "model": model,
                "dimension": len(result.data[0].embedding),
                "total_ms": total_ms,
                "embedding_preview": result.data[0].embedding[:8],
            })
        except Exception as exc:
            return JSONResponse({"error": str(exc)}, status_code=500)

    # ── CV Remix PRO Endpoints ──────────────────────────────────────

    @app.get("/cv")
    def cv_editor():
        cv_html = static_dir / "cv-remix-pro.html"
        if cv_html.is_file():
            return FileResponse(cv_html)
        return FileResponse(static_dir / "index.html")

    @app.post("/api/cv/parse")
    async def cv_parse(file: UploadFile = File(...)):
        filename = (file.filename or "upload").lower()
        contents = await file.read()
        if len(contents) > 20 * 1024 * 1024:
            raise HTTPException(400, "Arquivo muito grande. Maximo: 20MB")

        text = ""
        if filename.endswith(".pdf"):
            if PdfReader is None:
                raise HTTPException(500, "pypdf nao instalado. Execute: pip install pypdf")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            try:
                tmp.write(contents)
                tmp.close()
                reader = PdfReader(tmp.name)
                text = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            finally:
                os.unlink(tmp.name)
        elif filename.endswith(".docx"):
            if DocxDocument is None:
                raise HTTPException(500, "python-docx nao instalado. Execute: pip install python-docx")
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            try:
                tmp.write(contents)
                tmp.close()
                doc = DocxDocument(tmp.name)
                text = "\n".join(p.text for p in doc.paragraphs).strip()
            finally:
                os.unlink(tmp.name)
        elif filename.endswith(".txt"):
            text = contents.decode("utf-8", errors="replace").strip()
        else:
            raise HTTPException(400, "Formato nao suportado. Aceitos: PDF, DOCX, TXT")

        if len(text) < 20:
            raise HTTPException(400, "Nao foi possivel extrair texto do arquivo. Verifique se o documento contem texto selecionavel.")

        extraction_prompt = (
            "Analise o texto de curriculo abaixo e retorne APENAS um JSON valido (sem markdown, sem ```) "
            "com esta estrutura exata. NAO invente informacoes que nao existem no texto. "
            "Se algo nao estiver presente, deixe como string vazia ou array vazio.\n\n"
            "{\n"
            '  "personal": {"name":"","title":"","email":"","phone":"","location":"","country":"","nationality":"","dob":"","linkedin":"","portfolio":""},\n'
            '  "target_job": "",\n'
            '  "summary": "",\n'
            '  "experience": [{"company":"","role":"","period":"","responsibilities":[""],"results":[""]}],\n'
            '  "education": [{"institution":"","degree":"","period":"","details":""}],\n'
            '  "skills_technical": [""],\n'
            '  "skills_personal": [""],\n'
            '  "languages": [{"language":"","level":""}],\n'
            '  "certifications": [{"name":"","institution":"","year":"","validity":"","credential":""}],\n'
            '  "courses": [{"name":"","institution":"","period":""}],\n'
            '  "links": [{"label":"","url":""}],\n'
            '  "driving_license": {"has":false,"categories":[],"has_vehicle":false},\n'
            '  "availability": {"immediate":false,"shifts":false,"weekends":false,"relocations":false,"relocate":false},\n'
            '  "work_authorization": {"status":"","residence_permit":""}\n'
            "}\n\n"
            "REGRAS:\n"
            "- Se encontrar 'skills' ou 'habilidades' gerais, coloque em skills_technical\n"
            "- Se encontrar competencias pessoais (comunicacao, trabalho em equipa, etc), coloque em skills_personal\n"
            "- Se encontrar carta de conducao/categorias, preencha driving_license\n"
            "- Se encontrar disponibilidade, preencha availability\n"
            "- Se encontrar autorizacao de trabalho, preencha work_authorization\n"
            "- NAO invente informacoes que nao existem no texto original\n\n"
            f"Texto do curriculo ({filename}):\n\n{text[:12000]}"
        )

        cv_config = get_config("cv-remix-pro")
        cv_skill = cv_config.get("skill", "cv-remix-pro.md")
        skill_content = ""
        for sid, sdata in skills.items():
            if sid == "cv-remix-pro" or (sdata.get("label", "").lower().find("cv") >= 0):
                skill_content = sdata.get("content", "")
                break
        if not skill_content:
            skill_content = skills.get("cv-remix-pro", {}).get("content", "")

        client, params = client_for(DEFAULT_MODEL)
        messages = [
            {"role": "system", "content": skill_content or "Voce e um especialista em curriculos. Retorne apenas JSON valido."},
            {"role": "user", "content": extraction_prompt},
        ]

        try:
            completion = client.chat.completions.create(
                model=DEFAULT_MODEL, messages=messages, temperature=0.3, top_p=0.95, max_tokens=4096,
            )
            ai_response = completion.choices[0].message.content or "{}"
            ai_response = ai_response.strip()
            if ai_response.startswith("```"):
                ai_response = ai_response.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            cv_data = json.loads(ai_response)
        except json.JSONDecodeError:
            cv_data = {
                "personal": {"name": "", "title": "", "email": "", "phone": "", "location": "", "linkedin": "", "portfolio": ""},
                "summary": text[:500],
                "experience": [], "education": [], "skills": [], "languages": [],
                "certifications": [], "courses": [], "links": []
            }
        except Exception:
            cv_data = {
                "personal": {"name": "", "title": "", "email": "", "phone": "", "location": "", "linkedin": "", "portfolio": ""},
                "summary": text[:500],
                "experience": [], "education": [], "skills": [], "languages": [],
                "certifications": [], "courses": [], "links": []
            }

        return JSONResponse({"cv_data": cv_data, "raw_text": text[:3000]})

    @app.post("/api/cv/enhance")
    async def cv_enhance(request: Request):
        body = await request.json()
        cv_data = body.get("cv_data", {})
        enhancement_type = body.get("type", "summary")
        custom_prompt = body.get("prompt", "")
        target_job = body.get("target_job", "")
        language = body.get("language", "pt-pt")

        lang_context = {
            "pt-pt": "Portugues de Portugal. Use: curriculo, telemovel, carta de conducao, formacao profissional, competencias, funcao, experiencia profissional.",
            "pt-br": "Portugues do Brasil. Use: curriculo, celular, habilitacao, formacao, habilidades, cargo, experiencia profissional.",
            "en": "English professional curriculum.",
            "es": "Curriculum profesional en espanol.",
            "fr": "Curriculum professionnel en francais.",
            "de": "Professioneller Lebenslauf auf Deutsch.",
            "it": "Curriculum professionale in italiano."
        }

        prompts_map = {
            "summary": "Melhore o resumo profissional do curriculo. Mantenha o conteudo real. Apenas melhore a escrita, torne mais impactante e profissional. Nao invente informacoes novas. Retorne APENAS o curriculo atualizado no formato JSON.",
            "experience": "Melhore as experiencias profissionais do curriculo. Use verbos de acao, adicione metricas quando possivel, mantenha o conteudo real. Nao invente empresas ou cargos novos. Retorne APENAS o curriculo atualizado no formato JSON.",
            "grammar": "Corrija todos os erros ortograficos e gramaticais do curriculo. Nao altere o conteudo, apenas corrija a escrita. Retorne APENAS o curriculo atualizado no formato JSON.",
            "professional": "Reescreva o curriculo usando linguagem corporativa profissional, termos da industria e tom executivo. Mantenha todos os fatos inalterados. Retorne APENAS o curriculo atualizado no formato JSON.",
            "summarize": "Resuma o curriculo para caber em 1 pagina. Mantenha as informacoes mais importantes. Retorne APENAS o curriculo atualizado no formato JSON.",
            "expand": "Expanda as informacoes do curriculo, adicionando detalhes relevantes baseados no que ja existe. Nao invente experiencias novas. Retorne APENAS o curriculo atualizado no formato JSON.",
            "ats": "Otimize o curriculo para ATS (Applicant Tracking System). Adicione palavras-chave relevantes ao setor. Mantenha formatacao simples. Retorne APENAS o curriculo atualizado no formato JSON.",
            "adapt_pt": "Adapte o curriculo para o mercado de trabalho portugues. Use vocabulario profissional de Portugal: curriculo, telemovel, carta de conducao, formacao profissional, competencias, funcao, experiencia profissional. Evite brasileirismos. Nao invente informacoes. Retorne APENAS o curriculo atualizado no formato JSON.",
            "improve_all": "Melhore todo o curriculo: resumo, experiencia, competencias e formacao. Torne mais profissional e impactante. Nao invente informacoes. Retorne APENAS o curriculo atualizado no formato JSON.",
        }

        system_prompt = prompts_map.get(enhancement_type, prompts_map["summary"])
        if custom_prompt:
            system_prompt = custom_prompt

        lang_ctx = lang_context.get(language, lang_context["pt-pt"])
        system_prompt += f"\n\nIdioma do curriculo: {lang_ctx}"
        if target_job:
            system_prompt += f"\nO curriculo deve ser adaptado para o cargo: {target_job}. Nao invente experiencias ou competencias que o usuario nao possui."

        cv_text = json.dumps(cv_data, ensure_ascii=False, indent=2)

        client, params = client_for(DEFAULT_MODEL)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Curriculo atual:\n\n{cv_text[:10000]}"},
        ]

        return _sse_response(_sse_stream(client, DEFAULT_MODEL, messages, params))

    @app.post("/api/cv/translate")
    async def cv_translate(request: Request):
        body = await request.json()
        cv_data = body.get("cv_data", {})
        target_lang = body.get("target", "en")
        source_language = body.get("source_language", "pt-pt")

        lang_names = {
            "pt-pt": "portugues de Portugal profissional",
            "pt-br": "portugues do Brasil profissional",
            "en": "ingles profissional",
            "es": "espanhol profissional",
            "fr": "frances profissional",
            "de": "alemao profissional",
            "it": "italiano profissional"
        }
        lang_name = lang_names.get(target_lang, "ingles profissional")
        source_name = lang_names.get(source_language, "portugues de Portugal")

        preserve_rules = {
            "pt-pt": "Use vocabulario de Portugal: curriculo, telemovel, carta de conducao, formacao profissional, competencias, funcao, experiencia profissional.",
            "pt-br": "Use vocabulario do Brasil: curriculo, celular, habilitacao, formacao, habilidades, cargo, experiencia profissional.",
            "en": "Use professional English curriculum terminology.",
            "es": "Use terminologia profesional de curriculum en espanol.",
            "fr": "Use la terminologie professionnelle de curriculum en francais.",
            "de": "Verwenden Sie professionelle Lebenslauf-Terminologie auf Deutsch.",
            "it": "Usa la terminologia professionale del curriculum in italiano."
        }

        system_prompt = (
            f"Traduza este curriculo de {source_name} para {lang_name}. "
            f"{preserve_rules.get(target_lang, '')} "
            "Preserve: nomes proprios, empresas, datas, URLs, emails, numeros, certificacoes oficiais. "
            "O texto deve soar como um curriculo nativo. Nao faca traducao literal. "
            "Retorne APENAS o curriculo traduzido no mesmo formato JSON, incluindo todos os campos (personal, summary, experience, education, skills_technical, skills_personal, languages, certifications, courses, links, driving_license, availability, work_authorization)."
        )

        cv_text = json.dumps(cv_data, ensure_ascii=False, indent=2)

        client, params = client_for(DEFAULT_MODEL)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Curriculo para traduzir:\n\n{cv_text[:10000]}"},
        ]

        return _sse_response(_sse_stream(client, DEFAULT_MODEL, messages, params))

    @app.post("/api/cv/adapt")
    async def cv_adapt(request: Request):
        body = await request.json()
        cv_data = body.get("cv_data", {})
        job_description = body.get("job_description", "")
        job_url = body.get("job_url", "")
        job_title = body.get("job_title", "")
        language = body.get("language", "pt-pt")

        if job_url:
            if not job_url.startswith("http"):
                raise HTTPException(400, "URL deve comecar com http:// ou https://")
            try:
                page_data = _fetch_and_clean_url(job_url)
                job_description = f"Titulo: {page_data['title']}\n\n{page_data['content'][:5000]}"
            except Exception:
                pass

        if not job_description and not job_title:
            raise HTTPException(400, "Cargo, descricao da vaga ou URL nao informado.")

        lang_rules = {
            "pt-pt": "Portugues de Portugal. Use: curriculo, telemovel, carta de conducao, formacao profissional, competencias, funcao, experiencia profissional.",
            "pt-br": "Portugues do Brasil. Use: curriculo, celular, habilitacao, formacao, habilidades, cargo, experiencia profissional.",
            "en": "English.", "es": "Espanol.", "fr": "Francais.", "de": "Deutsch.", "it": "Italiano."
        }

        system_prompt = (
            f"Adapte o curriculo para o cargo: {job_title or 'conforme descricao da vaga'}. "
            f"Idioma: {lang_rules.get(language, lang_rules['pt-pt'])} "
            "REGRAS ABSOLUTAS:\n"
            "- PODE: melhorar escrita, destacar experiencias relevantes, reorganizar competencias, adaptar resumo profissional, alterar ordem das informacoes, melhorar vocabulario profissional.\n"
            "- NAO PODE: inventar experiencia, inventar formacao, inventar certificacoes, inventar empresas, inventar habilitacoes, inventar carta de conducao, inventar datas.\n"
            "Retorne APENAS o curriculo adaptado no mesmo formato JSON, incluindo todos os campos (personal, target_job, summary, experience, education, skills_technical, skills_personal, languages, certifications, courses, links, driving_license, availability, work_authorization)."
        )

        cv_text = json.dumps(cv_data, ensure_ascii=False, indent=2)

        user_content = f"CURRICULO:\n\n{cv_text[:8000]}"
        if job_title:
            user_content += f"\n\n---\n\nCARGO DESEJADO: {job_title}"
        if job_description:
            user_content += f"\n\n---\n\nDESCRICAO DA VAGA:\n\n{job_description[:5000]}"

        client, params = client_for(DEFAULT_MODEL)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content},
        ]

        return _sse_response(_sse_stream(client, DEFAULT_MODEL, messages, params))

    @app.post("/api/cv/export-docx")
    async def cv_export_docx(request: Request):
        if DocxDocument is None:
            raise HTTPException(500, "python-docx nao instalado. Execute: pip install python-docx")

        body = await request.json()
        cv_data = body.get("cv_data", {})
        template = body.get("template", "europe")

        doc = DocxDocument()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        personal = cv_data.get("personal", {})
        name = personal.get("name", "Curriculo")

        if personal.get("name"):
            heading = doc.add_heading(personal["name"], level=0)
            for run in heading.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        if personal.get("location"):
            contact_parts.append(personal["location"])
        if personal.get("country"):
            contact_parts.append(personal["country"])
        if personal.get("nationality"):
            contact_parts.append("Nac.: " + personal["nationality"])
        if contact_parts:
            p = doc.add_paragraph(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if personal.get("linkedin"):
            p = doc.add_paragraph("LinkedIn: " + personal["linkedin"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)

        if personal.get("title"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(personal["title"])
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        target_job = cv_data.get("target_job", "")
        if target_job:
            p = doc.add_paragraph()
            run = p.add_run("Cargo pretendido: " + target_job)
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

        if cv_data.get("summary"):
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(cv_data["summary"])

        if cv_data.get("experience"):
            doc.add_heading("Work Experience", level=1)
            for exp in cv_data["experience"]:
                if exp.get("role") or exp.get("company"):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}")
                    run.bold = True
                    if exp.get("period"):
                        run2 = p.add_run(f"  ({exp['period']})")
                        run2.font.size = Pt(9)
                        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                for resp in exp.get("responsibilities", []):
                    if resp:
                        doc.add_paragraph(resp, style="List Bullet")
                for res in exp.get("results", []):
                    if res:
                        doc.add_paragraph(res, style="List Bullet")

        if cv_data.get("education"):
            doc.add_heading("Education", level=1)
            for edu in cv_data["education"]:
                if edu.get("degree") or edu.get("institution"):
                    p = doc.add_paragraph()
                    run = p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
                    run.bold = True
                    if edu.get("period"):
                        run2 = p.add_run(f"  ({edu['period']})")
                        run2.font.size = Pt(9)
                if edu.get("details"):
                    doc.add_paragraph(edu["details"])

        skills_tech = cv_data.get("skills_technical", cv_data.get("skills", []))
        skills_pers = cv_data.get("skills_personal", [])
        if skills_tech and any(s for s in skills_tech if s):
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(" | ".join(s for s in skills_tech if s))
        if skills_pers and any(s for s in skills_pers if s):
            doc.add_heading("Personal Skills", level=1)
            doc.add_paragraph(" | ".join(s for s in skills_pers if s))

        if cv_data.get("languages"):
            doc.add_heading("Languages", level=1)
            for lang in cv_data["languages"]:
                if lang.get("language"):
                    doc.add_paragraph(f"{lang['language']} - {lang.get('level', '')}")

        certs = cv_data.get("certifications", [])
        certs = [c for c in certs if c.get("name")]
        if certs:
            doc.add_heading("Certifications", level=1)
            for cert in certs:
                parts = [cert["name"]]
                if cert.get("institution"):
                    parts.append(cert["institution"])
                if cert.get("year"):
                    parts.append(cert["year"])
                if cert.get("validity"):
                    parts.append("Val: " + cert["validity"])
                if cert.get("credential"):
                    parts.append("No: " + cert["credential"])
                doc.add_paragraph(" - ".join(parts))

        if cv_data.get("courses"):
            courses = [c for c in cv_data["courses"] if c.get("name")]
            if courses:
                doc.add_heading("Courses", level=1)
                for course in courses:
                    parts = [course["name"]]
                    if course.get("institution"):
                        parts.append(course["institution"])
                    if course.get("period"):
                        parts.append(course["period"])
                    doc.add_paragraph(" - ".join(parts))

        dl = cv_data.get("driving_license", {})
        if dl.get("has"):
            cats = [c for c in dl.get("categories", []) if c]
            driving_parts = []
            if cats:
                driving_parts.append("Categories: " + ", ".join(cats))
            if dl.get("has_vehicle"):
                driving_parts.append("Own vehicle")
            if driving_parts:
                doc.add_heading("Driving License", level=1)
                doc.add_paragraph(" | ".join(driving_parts))

        avail = cv_data.get("availability", {})
        avail_items = []
        if avail.get("immediate"): avail_items.append("Immediate availability")
        if avail.get("shifts"): avail_items.append("Shifts")
        if avail.get("weekends"): avail_items.append("Weekends")
        if avail.get("relocations"): avail_items.append("Relocations")
        if avail.get("relocate"): avail_items.append("Willing to relocate")
        if avail_items:
            doc.add_heading("Availability", level=1)
            doc.add_paragraph(" | ".join(avail_items))

        wa = cv_data.get("work_authorization", {})
        if wa.get("status") or wa.get("residence_permit"):
            doc.add_heading("Work Authorization", level=1)
            auth_parts = []
            if wa.get("status"): auth_parts.append("Authorization: " + wa["status"])
            if wa.get("residence_permit"): auth_parts.append(wa["residence_permit"])
            doc.add_paragraph(" | ".join(auth_parts))

        safe_name = re.sub(r"[^\w\s\-]", "", name or "Curriculo").strip().replace(" ", "-")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        try:
            doc.save(tmp.name)
            tmp.close()
            return FileResponse(
                tmp.name,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{safe_name}-CV.docx",
                background=None,
            )
        except Exception:
            if os.path.exists(tmp.name):
                os.unlink(tmp.name)
            raise HTTPException(500, "Erro ao gerar DOCX")

    @app.post("/api/cv/export-pdf")
    async def cv_export_pdf(request: Request):
        body = await request.json()
        cv_data = body.get("cv_data", {})
        template = body.get("template", "europe")

        html = _generate_cv_html(cv_data, template)
        return JSONResponse({"html": html})

    return app
